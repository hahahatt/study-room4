import pandas as pd
from io import BytesIO
from docx import Document
from .pii_masker import mask_text

def _seek0(f):
    try:
        f.seek(0)
    except Exception:
        pass

def create_bytes_io(text, original_filename):
    output = BytesIO()
    output.write(text.encode("utf-8"))
    output.name = "masked_" + original_filename
    output.seek(0)
    return output

def mask_attachment(file_obj):
    filename = file_obj.name
    ext = filename.split('.')[-1].lower()

    try:
        if ext == "txt":
            _seek0(file_obj)
            text = file_obj.read().decode("utf-8", errors="ignore")
            masked = mask_text(text)
            return create_bytes_io(masked, filename)

        elif ext == "csv":
            _seek0(file_obj)
            try:
                df = pd.read_csv(file_obj, encoding="utf-8-sig")
            except UnicodeDecodeError:
                _seek0(file_obj)
                df = pd.read_csv(file_obj, encoding="euc-kr")

            for col in df.columns:
                df[col] = df[col].astype(str).apply(mask_text)

            output = BytesIO()
            df.to_csv(output, index=False, encoding="utf-8-sig")
            output.name = "masked_" + filename
            output.seek(0)
            return output

        elif ext == "xlsx":
            from openpyxl import load_workbook
            _seek0(file_obj)
            wb = load_workbook(file_obj)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str):
                            cell.value = mask_text(cell.value)

            output = BytesIO()
            wb.save(output)
            output.name = "masked_" + filename
            output.seek(0)
            return output

        elif ext == "docx":
            _seek0(file_obj)
            doc = Document(file_obj)

            for para in doc.paragraphs:
                para.text = mask_text(para.text)

            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        original = cell.text or ""
                        masked = mask_text(original)
                        if masked != original:
                            cell.text = masked

            output = BytesIO()
            doc.save(output)
            output.name = "masked_" + filename
            output.seek(0)
            return output

        elif ext == "pdf":
            try:
                import fitz  # PyMuPDF
                from PIL import Image, ImageDraw
                import pytesseract
                import re

                _seek0(file_obj)
                pdf_bytes = file_obj.read()
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")

                def is_digital_pdf(doc):
                    total = 0
                    for page in doc:
                        total += len((page.get_text("text") or "").strip())
                        if total > 100:
                            break
                    return total > 100

                patterns = [
                    re.compile(r"\d{6}-\d{7}"),
                    re.compile(r"01[016789]-?\d{3,4}-?\d{4}"),
                    re.compile(r"\b[\w\.-]+@[\w\.-]+\.\w{2,}\b"),
                    re.compile(r"\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}"),
                    re.compile(r"(기밀|내부문서|confidential|비공개|internal use only)", re.I)
                ]

                if is_digital_pdf(doc):
                    for page in doc:
                        text = page.get_text("text")
                        targets = set()
                        for pat in patterns:
                            targets.update(m.group(0) for m in pat.finditer(text))
                        for s in targets:
                            for r in page.search_for(s):
                                page.add_redact_annot(r, fill=(0, 0, 0))

                        for x0, y0, x1, y1, word, *_ in page.get_text("words"):
                            try:
                                if mask_text(word) != word:
                                    page.add_redact_annot(fitz.Rect(x0, y0, x1, y1), fill=(0, 0, 0))
                            except Exception:
                                pass

                        page.apply_redactions()

                    output = BytesIO()
                    doc.save(output)
                    output.name = "masked_" + filename
                    output.seek(0)
                    doc.close()
                    return output

                else:
                    dpi = 250
                    zoom = dpi / 72.0
                    mat = fitz.Matrix(zoom, zoom)
                    out_doc = fitz.open()

                    for page in doc:
                        pix = page.get_pixmap(matrix=mat, alpha=False)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        draw = ImageDraw.Draw(img)

                        data = pytesseract.image_to_data(img, lang="kor+eng", output_type=pytesseract.Output.DICT)

                        for i in range(len(data["text"])):
                            token = (data["text"][i] or "").strip()
                            if not token:
                                continue

                            regex_hit = any(re.search(pat, token) for pat in [
                                r"\d{6}-\d{7}",
                                r"01[016789]-?\d{3,4}-?\d{4}",
                                r"\b[\w\.-]+@[\w\.-]+\.\w{2,}\b",
                                r"\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}",
                                r"(기밀|내부문서|confidential|비공개|internal use only)"
                            ])

                            ner_hit = False
                            try:
                                if mask_text(token) != token:
                                    ner_hit = True
                            except Exception:
                                pass

                            if regex_hit or ner_hit:
                                x, y, w, h = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
                                draw.rectangle([x, y, x + w, y + h], fill=(0, 0, 0))

                        img_bytes = BytesIO()
                        img.save(img_bytes, format="PNG")
                        img_bytes.seek(0)

                        rect = fitz.Rect(0, 0, page.rect.width, page.rect.height)
                        new_page = out_doc.new_page(width=page.rect.width, height=page.rect.height)
                        new_page.insert_image(rect, stream=img_bytes.read())

                    out = BytesIO()
                    out_doc.save(out)
                    out_doc.close()
                    out.name = "masked_" + filename
                    out.seek(0)
                    doc.close()
                    return out

            except Exception as e:
                print(f"[ERROR] PDF 마스킹 실패: {e}")
                return None
        elif ext in ["jpg", "png"]:
            from PIL import Image, ImageDraw
            import re
            import pytesseract

            pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'


            _seek0(file_obj)
            img = Image.open(file_obj)
            draw = ImageDraw.Draw(img)

            data = pytesseract.image_to_data(img, lang="kor+eng", output_type=pytesseract.Output.DICT)

            for i, text in enumerate(data["text"]):
                token = (text or "").strip()
                if not token:
                    continue

                # regex = any(re.search(pat, token) for pat in [
                #     r"\d{6}-\d{7}",
                #     r"010-?\d{4}=?\d{4}",
                #     r"\b[\w\.-]+@['\w\.-]+\.\w{2,}\b",
                #     r"\d{4}{-\s}?\d{4}[-\s]?\d{4}[-\s]?\d{4}"
                # ])

                ner_state = False
                if mask_text(token) != token:
                    ner_state = True

                #if regex or ner_state:
                if ner_state:
                    x, y, w, h = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
                    draw.rectangle([x,y, x+w, y+h], fill='black')
            
            output = BytesIO()
            img.save(output, format='png')
            output.name = "masked_"+filename
            output.seek(0)
            

            return output

        else:
            return None

    except Exception as e:
        print(f"[ERROR] 첨부파일 마스킹 실패: {e}")
        return None
