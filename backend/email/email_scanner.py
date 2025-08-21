from .pii_masker import mask_text, contains_sensitive_keywords, ner_predict, merge_entities
import pandas as pd
import docx
import PyPDF2
import pytesseract
from io import BytesIO
from PIL import Image
import re

# pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# ---------------- 스캔 요약본 시작 ----------------
# NER 라벨 기반 항목 표준화
_PII_LABELS = ["이름", "주민번호", "전화번호", "이메일", "카드번호", "주소", "계좌번호"]

# 계좌번호는 NER 라벨에 없으므로 정규식으로만 카운트(국내 포맷 중심)
_ACCOUNT_REGEXES = [
    re.compile(r'\b\d{2,4}-\d{2,4}-\d{2,6}\b'),
    re.compile(r'\b\d{10,14}\b'),  # 하이픈 없는 10~14자리
]

def _seek0(f):
    try: f.seek(0)
    except: pass

def _read_text_from_any(file):
    name = getattr(file, "name", "") or ""
    lower = name.lower()
    _seek0(file)
    raw = file.read()
    _seek0(file)

    try:
        if lower.endswith(".txt"):
            return raw.decode("utf-8", "ignore")
        if lower.endswith(".csv"):
            _seek0(file); import pandas as pd
            df = pd.read_csv(file, dtype=str, encoding="utf-8", engine="python")
            return df.to_string(index=False)
        if lower.endswith(".xlsx"):
            _seek0(file); import pandas as pd
            df = pd.read_excel(file, dtype=str, engine="openpyxl")
            return df.astype(str).to_string(index=False)
        if lower.endswith(".docx"):
            _seek0(file); d = docx.Document(file)
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            return "\n".join(parts)
        if lower.endswith(".pdf"):
            # 1) 디지털 PDF: 추출
            _seek0(file)
            text = []
            try:
                reader = PyPDF2.PdfReader(file)
                for p in reader.pages:
                    text.append(p.extract_text() or "")
                text = "\n".join(text).strip()
            except Exception:
                text = ""
            # 2) 비어있으면 OCR 폴백
            if not text:
                _seek0(file)
                import fitz
                doc = fitz.open(stream=file.read(), filetype="pdf")
                ocr_chunks = []
                for page in doc:
                    pix = page.get_pixmap(dpi=250)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_chunks.append(pytesseract.image_to_string(img, lang="kor+eng"))
                text = "\n".join(ocr_chunks)
            return text
        if lower.endswith((".jpg",".jpeg",".png")):
            _seek0(file)
            img = Image.open(BytesIO(raw))
            return pytesseract.image_to_string(img, lang="kor+eng")
    except Exception:
        pass
    return ""

def _count_with_ner_and_regex(text: str):
    # 기본 0으로 초기화
    counts = {k: 0 for k in _PII_LABELS}
    if not text:
        return counts

    # 1) NER로 추출 → 이름/주민/전화/이메일/카드/주소 카운트
    try:
        tagged = ner_predict(text)               # [('홍', 'B-이름'), ...] 형태
        entities = merge_entities(tagged)        # [('홍길동','이름'), ...]
        for _, label in entities:
            if label in counts and label != "계좌번호":
                counts[label] += 1
    except Exception:
        pass

    # 2) 계좌번호는 정규식
    acc_total = 0
    for rgx in _ACCOUNT_REGEXES:
        acc_total += len(rgx.findall(text))
    counts["계좌번호"] = max(counts["계좌번호"], acc_total)

    # 3) 보수적 보강(정규식으로 더 잘 잡히는 것들은 max로 상향)
    #    - 중복 카운트를 피하려고 합이 아닌 max 사용
    email_re = re.compile(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}')
    phone_re = re.compile(r'(?<!\d)(?:01[016789]|02|0[3-6][1-5])[-.\s]?\d{3,4}[-.\s]?\d{4}(?!\d)')
    card_re  = re.compile(r'(?<!\d)(?:\d{4}[-\s]?){3}\d{4}(?!\d)')
    jumin_re = re.compile(r'(?<!\d)\d{6}[-.\s]?\d{7}(?!\d)')
    addr_re  = re.compile(r'(?:[가-힣]{2,}(?:시|도)\s*)?[가-힣]{2,}(?:구|군)\s+[가-힣0-9\-]{1,}(?:로|길)\s*\d{1,4}')

    counts["이메일"]   = max(counts["이메일"],   len(email_re.findall(text)))
    counts["전화번호"] = max(counts["전화번호"], len(phone_re.findall(text)))
    counts["카드번호"] = max(counts["카드번호"], len(card_re.findall(text)))
    counts["주민번호"] = max(counts["주민번호"], len(jumin_re.findall(text)))
    counts["주소"]     = max(counts["주소"],     len(addr_re.findall(text)))

    return counts

def summarize_email(subject: str, body: str, attachments) -> pd.DataFrame:
    """스캔 요약본 표(DataFrame) 생성 — ✅ 첨부만 집계(본문 제외)"""
    agg = {k: 0 for k in _PII_LABELS}
    for f in (attachments or []):
        try:
            text = _read_text_from_any(f)
        except Exception:
            text = ""
        c = _count_with_ner_and_regex(text)
        for k in _PII_LABELS:
            agg[k] += c[k]

    rows = [{"항목": k, "첨부": agg[k]} for k in _PII_LABELS]
    df = pd.DataFrame(rows, columns=["항목", "첨부"])
    return df
# ---------------- 스캔 요약본 끝 ----------------










def scan_image(file_obj):
    """
    이미지 파일에서 텍스트를 추출하여 문자열로 반환하는 함수.
    """
    try:
        print('===============================')
        print('Start!!\n', file_obj)
        img = Image.open(file_obj)
        print('===============================')
        
        text = pytesseract.image_to_string(img, lang='kor+eng')
        
        return text
    except FileNotFoundError:
        return "오류: 이미지를 찾을 수 없습니다. 경로를 다시 확인해 주세요."
    except Exception as e:
        return f"텍스트 추출 중 오류가 발생했습니다: {e}"

# 📄 OCR 기반 PDF 텍스트 추출 함수
def extract_text_from_pdf_with_ocr(file_obj):
    import pytesseract
    import fitz  # PyMuPDF
    from PIL import Image

    file_obj.seek(0)
    doc = fitz.open(stream=file_obj.read(), filetype="pdf")
    full_text = ""

    for page in doc:
        pix = page.get_pixmap(dpi=250)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img, lang="kor+eng")
        full_text += text + "\n"

    return full_text

# 📬 메일 본문/제목/첨부파일 민감정보 스캐너
def scan_email(subject, body, attachments):
    warnings = []
    masked_body = mask_text(body)

    # 제목 키워드 탐지
    if contains_sensitive_keywords(subject, ["기밀", "내부문서", "계약서", "극비", "대외비"]):
        warnings.append("제목에 기밀 관련 키워드 포함")

    # 본문 마스킹 여부 확인
    if masked_body != body:
        warnings.append("본문에 개인정보 또는 민감 정보 포함")

    # 첨부파일 검사
    for file in attachments or []:
        try:
            file.seek(0)
        except Exception:
            pass

        filename = file.name.lower()
        content = ""

        try:
            if filename.endswith(".txt"):
                file.seek(0)
                content = file.read().decode("utf-8", errors="ignore")

            elif filename.endswith(".csv"):
                file.seek(0)
                df = pd.read_csv(file)
                content = df.to_string()

            elif filename.endswith(".xlsx"):
                file.seek(0)
                df = pd.read_excel(file)
                content = df.to_string()

            elif filename.endswith(".docx"):
                file.seek(0)
                docx_obj = docx.Document(file)
                content_parts = []

                # 본문
                content_parts.extend([p.text for p in docx_obj.paragraphs if p.text.strip()])

                # 테이블
                for table in docx_obj.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content_parts.append(cell.text)

                content = "\n".join(content_parts)

            elif filename.endswith(".pdf"):
                file.seek(0)
                reader = PyPDF2.PdfReader(file)
                content = "\n".join([page.extract_text() or "" for page in reader.pages])

                if not content.strip():
                    content = extract_text_from_pdf_with_ocr(file)

            elif filename.endswith(".jpg") or filename.endswith(".png"):
                file.seek(0)
                img_bytes = BytesIO(file.read())
                content = scan_image(img_bytes)

            else:
                continue

        except Exception as e:
            warnings.append(f"첨부파일 `{file.name}` 처리 중 오류 발생: {e}")
            continue

        # 민감 키워드 또는 마스킹 전후 비교로 탐지
        if (
            contains_sensitive_keywords(content, ["기밀", "내부자료", "주민번호", "카드번호", "주소", "전화번호", "이메일", "계좌번호", "이름"]) or
            mask_text(content) != content
        ):
            warnings.append(f"첨부파일 `{file.name}` 에 개인정보 또는 민감 정보 포함")

    return warnings, masked_body
